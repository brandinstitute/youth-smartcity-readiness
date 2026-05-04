"""Repair the monograph after Word delete-only edits.

This is a one-off forensic recovery script kept in the repository for
reproducibility. It expects a partially-broken Word draft at
data/monograph_in.docx and writes the repaired version to
results/monograph_repaired.docx. Each fix is annotated below with the line
of damage it targets — review and adapt before running on a different file.

Strategy: open the .docx, walk paragraphs and tables, fix specific known-broken
spots by REPLACING entire paragraph runs (we already know exact context lines).
This is safer than trusting fuzzy text-search inside Word.
"""
from pathlib import Path
from docx import Document
from copy import deepcopy

ROOT = Path(__file__).resolve().parents[1]
SRC = str(ROOT / "data" / "monograph_in.docx")
DST = str(ROOT / "results" / "monograph_repaired.docx")
(Path(DST).parent).mkdir(parents=True, exist_ok=True)

doc = Document(SRC)


def set_paragraph_text(para, new_text):
    """Replace the paragraph's visible text with `new_text`, preserving the first run's formatting."""
    runs = para.runs
    if not runs:
        # Add a fresh run
        run = para.add_run(new_text)
        return
    # Set first run to new text, clear remaining runs' text
    runs[0].text = new_text
    for r in runs[1:]:
        r.text = ''


def find_paragraph_by_prefix(prefix, start=0):
    """Find first paragraph whose text starts with `prefix`. Return index or -1."""
    for i, p in enumerate(doc.paragraphs[start:], start=start):
        if p.text.startswith(prefix):
            return i
    return -1


def find_paragraph_by_substring(needle, start=0):
    for i, p in enumerate(doc.paragraphs[start:], start=start):
        if needle in p.text:
            return i
    return -1


print(f'Document loaded. Paragraphs: {len(doc.paragraphs)}, tables: {len(doc.tables)}')

# ============================================================
# FIX 1 — §4.4.1 AEI input variables (riadky ~482-483)
# ============================================================
i = find_paragraph_by_substring('POCETAPLIK – number of')
if i >= 0:
    set_paragraph_text(doc.paragraphs[i],
        'POCETAPLIK – self-reported number of installed applications (portfolio size).')
    print(f'  FIX 1a: paragraph #{i} -> POCETAPLIK fixed')
else:
    print('  FIX 1a: NOT FOUND')

i = find_paragraph_by_substring('recency_new_app – recency of last new app installation ()')
if i >= 0:
    set_paragraph_text(doc.paragraphs[i],
        'recency_new_app – recency of last new app installation (1–3 ordinal: 1 = more than a month ago, 2 = within the last month, 3 = within the last week).')
    print(f'  FIX 1b: paragraph #{i} -> recency fixed')
else:
    print('  FIX 1b: NOT FOUND')

# Rationale broken bullet
i = find_paragraph_by_substring('number of ed apps (breadth of the app ),')
if i >= 0:
    set_paragraph_text(doc.paragraphs[i],
        'number of installed apps (breadth of the app portfolio that the student maintains), and')
    print(f'  FIX 1c: paragraph #{i} -> rationale bullet fixed')
else:
    print('  FIX 1c: NOT FOUND')

# ============================================================
# FIX 12 — §4.4.1 recency mapping (riadky ~328-332)
# ============================================================
i = find_paragraph_by_substring('Recoded to a –3 scale')
if i >= 0:
    set_paragraph_text(doc.paragraphs[i],
        'Recoded to a 1–3 ordinal scale:')
    print(f'  FIX 12a: paragraph #{i} -> "Recoded" header fixed')
    # The next 4 paragraphs are empty placeholders. Fill them.
    targets = [
        '1 = "More than a month ago"',
        '2 = "Within the last month"',
        '3 = "Within the last week (including today)"',
    ]
    filled = 0
    for j in range(i+1, min(i+8, len(doc.paragraphs))):
        if doc.paragraphs[j].text.strip() == '' and filled < len(targets):
            set_paragraph_text(doc.paragraphs[j], targets[filled])
            filled += 1
            print(f'  FIX 12b: paragraph #{j} <- "{targets[filled-1]}"')
        elif doc.paragraphs[j].text.strip() != '':
            break
else:
    print('  FIX 12: NOT FOUND')

# ============================================================
# FIX 6 — §6.3.2 "Of the 3 selected the conceptually correct..."
# ============================================================
i = find_paragraph_by_substring('Of the 3 selected the conceptually correct')
if i >= 0:
    old = doc.paragraphs[i].text
    new = old.replace(
        'Of the 3 selected the conceptually correct answer',
        'Of the 368 students in the analytical sample, 342 (92.9%) selected the conceptually correct answer')
    set_paragraph_text(doc.paragraphs[i], new)
    print(f'  FIX 6: paragraph #{i} -> Of the 368 ... 342 (92.9%)')
else:
    print('  FIX 6: NOT FOUND')

# ============================================================
# FIX 8 — §4.4.5 mean-imputation note  (insert after the rationale paragraph)
# ============================================================
sentence_8 = ('Missing-data handling. Where a respondent provided a valid CSA score but had a missing '
              'component for SDR or DAE, the missing first-order index was replaced by its column mean '
              'before applying the weighted formula. This light imputation preserves valid N = 368 and '
              'at most ~13% of cases are affected; a fully listwise sensitivity variant '
              '(YSCR_listwise, valid N ≈ 332) is computed in parallel and correlates with the default '
              'at Pearson r ≈ 0.99, confirming that the substantive ranking of students is unchanged.')

# Find the YSCR rationale paragraph (the one ending with "comparative metric ... etc.")
i = find_paragraph_by_substring('not to create an absolute, universal benchmark')
if i >= 0:
    # Insert new paragraph after this one
    para = doc.paragraphs[i]
    new_para = deepcopy(para._element)
    para._element.addnext(new_para)
    new = doc.paragraphs[i+1]
    set_paragraph_text(new, sentence_8)
    print(f'  FIX 8: inserted after paragraph #{i}')
else:
    print('  FIX 8: NOT FOUND')

# ============================================================
# FIX 11 — §4.1.1 (Dataset A sample) 83 → 82 explanation
# ============================================================
sentence_11 = ('The LimeSurvey export contains 83 secondary-school respondents; one case lacking a '
               'valid gender response was excluded from the analytical sample, yielding the working '
               'N = 82 used throughout Chapters 5–8.')

# Find the line "Number of respondents: 82 valid completed questionnaires from secondary-school..."
i = find_paragraph_by_substring('Number of respondents: 82 valid completed questionnaires')
if i >= 0:
    para = doc.paragraphs[i]
    new_para = deepcopy(para._element)
    para._element.addnext(new_para)
    new = doc.paragraphs[i+1]
    set_paragraph_text(new, sentence_11)
    print(f'  FIX 11: inserted after paragraph #{i}')
else:
    print('  FIX 11: NOT FOUND')

# ============================================================
# Now fix tables: Tab 13, Tab 32, Tab 35, Tab 40
# ============================================================

def set_cell_text(cell, text):
    """Replace cell text by clearing the first paragraph's runs and writing text."""
    # Wipe additional paragraphs in cell (keep first)
    for p in cell.paragraphs[1:]:
        p._element.getparent().remove(p._element)
    p = cell.paragraphs[0]
    set_paragraph_text(p, text)


print(f'\nTotal tables in document: {len(doc.tables)}')
for ti, t in enumerate(doc.tables):
    first = t.rows[0].cells[0].text[:50] if t.rows else ''
    print(f'  Table {ti}: rows={len(t.rows)}, cols={len(t.rows[0].cells) if t.rows else 0}, first cell="{first}"')


# ============================================================
# FIX 7 — Tab 13 (Internet/devices descriptives) — fill empty first cell
# ============================================================
# Locate the table where row 1 col 1 is empty but col 2 = '358'
tab13 = None
for ti, t in enumerate(doc.tables):
    for row in t.rows:
        cells = row.cells
        if len(cells) >= 2 and cells[0].text.strip() == '' and cells[1].text.strip() == '358':
            tab13 = (ti, row, cells)
            break
    if tab13:
        break
if tab13:
    ti, row, cells = tab13
    set_cell_text(cells[0], 'Mean devices per category (PC / tablet / smartphone / other)')
    print(f'  FIX 7: Table {ti} devices row label set')
else:
    print('  FIX 7: NOT FOUND')

# ============================================================
# FIX 2 — Tab 32 (AEI construction) row 1 Details cell
# ============================================================
for ti, t in enumerate(doc.tables):
    if len(t.rows) >= 5 and t.rows[0].cells[0].text.strip() == 'Step':
        # Check whether row 4 says "Min–max transform of AEI_raw to AEI"
        if any('AEI_raw to AEI' in c.text for c in t.rows[4].cells):
            # This is Tab 32
            set_cell_text(t.rows[1].cells[2], 'time_apps_h, apps_installed, recency_new_app')
            print(f'  FIX 2: Table {ti} (AEI construction) row 1 Details set')
            # Also fix the standardisation row (row 3 currently says "apps_used_z")
            std_text = 'Compute time_apps_h_z, apps_installed_z, recency_new_app_z (z-scores)'
            set_cell_text(t.rows[3].cells[2], std_text)
            print(f'  FIX 2b: standardisation row updated to apps_installed_z')
            break

# ============================================================
# FIX 4 — Tab 35 (CSA construction) — repair empty cells
# ============================================================
for ti, t in enumerate(doc.tables):
    if len(t.rows) >= 4 and t.rows[0].cells[0].text.strip() == 'Step':
        if any('CSA_raw to CSA' in c.text for c in t.rows[-1].cells):
            # Tab 35
            inputs = ('school_exposure (3-level: 0/1/2 from G01Q22), '
                      'role_correct (binary from G01Q10), '
                      'smart_edu_importance (Likert 1–5 from G01Q23), '
                      'local_features (binary from G05Q18)')
            std = 'Compute z-scores of all four items: z_school_exposure, z_role_correct, z_smart_edu_importance, z_local_features'
            agg = 'Compute CSA_raw as the mean of the four z-scores; respondents with at least three valid items are retained'
            set_cell_text(t.rows[1].cells[2], inputs)
            set_cell_text(t.rows[2].cells[2], std)
            set_cell_text(t.rows[3].cells[2], agg)
            print(f'  FIX 4: Table {ti} (CSA construction) repaired')
            break

# ============================================================
# FIX 9 — Tab 40 (segments) — fill all 4 segment rows
# ============================================================
segments = [
    ('Pragmatic utilitarians', '31.2% (N=24)', '40.95',
     '65–70', '49–52',
     'Moderate AEI, broad motives including school/work and curiosity; balanced app sets'),
    ('Social-entertainment heavy users', '32.5% (N=25)', '57.05',
     '65–70', '49–52',
     'High AEI, narrow motive set dominated by entertainment; frequent social/streaming use'),
    ('Low-engagement / reluctant', '22.1% (N=17)', '37.74',
     '≈ sample mean', '≈ sample mean',
     'Lower AEI, low novelty-seeking, narrow motives'),
    ('Trend followers / early adopters', '14.3% (N=11)', '63.02',
     '≈ sample mean', '≈ sample mean',
     'Highest AEI; multi-motive (entertainment, curiosity, trends); responsive to advertising'),
]

for ti, t in enumerate(doc.tables):
    if len(t.rows) >= 5:
        header_text = ' '.join(c.text for c in t.rows[0].cells)
        if 'Share of sample A' in header_text and 'AEI (mean)' in header_text:
            # This is Tab 40
            for ri, segment in enumerate(segments, start=1):
                if ri >= len(t.rows):
                    break
                row = t.rows[ri]
                for ci, val in enumerate(segment):
                    if ci < len(row.cells):
                        set_cell_text(row.cells[ci], val)
            print(f'  FIX 9: Table {ti} (segments) populated with {len(segments)} rows')
            break

# ============================================================
# FIX 10 — Drobné typo opravy
# ============================================================
# Glossary: "high school z students" → "high school students"
for i, p in enumerate(doc.paragraphs):
    if 'Ch. 5 (Digital habits of high school z' in p.text:
        new = p.text.replace('high school z', 'high school students (results A))').replace('students (results A))', 'students (results A))')
        # Actually the glossary entry likely splits across paragraphs; do a clean replace
        new = p.text.replace('high school z\n', 'high school ').replace('high school z', 'high school')
        set_paragraph_text(p, new)
        print(f'  FIX 10a: paragraph #{i} glossary "high school z" cleaned')

# Heading "Model :" -> "Model 7.1:" / "Model 7.2:"
for i, p in enumerate(doc.paragraphs):
    if p.style.name.startswith('Heading') and p.text.startswith('Model :'):
        # Determine context: first occurrence is 7.1 (Dataset A), second is 7.2 (Dataset B)
        if 'Dataset A' in p.text or 'App Engagement' in p.text:
            new = p.text.replace('Model :', 'Model 7.1:')
        elif 'Dataset B' in p.text or 'Awareness' in p.text:
            new = p.text.replace('Model :', 'Model 7.2:')
        else:
            new = p.text.replace('Model :', 'Model 7.1:')
        set_paragraph_text(p, new)
        print(f'  FIX 10b: paragraph #{i} -> {new[:60]}')

# Duplicate "respect privacy that respect privacy"
for i, p in enumerate(doc.paragraphs):
    if 'respect privacy that respect privacy' in p.text:
        new = p.text.replace('respect privacy that respect privacy', 'respect privacy')
        set_paragraph_text(p, new)
        print(f'  FIX 10d: paragraph #{i} duplicated "respect privacy" cleaned')

# ============================================================
doc.save(DST)
print(f'\n=== Saved to {DST} ===')
