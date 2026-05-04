"""Insert Appendix A (questionnaires) into the monograph at the existing
'Appendix A' placeholder line.

Strategy: locate the paragraph 'Normal: Appendix A' (currently empty body)
and append rendered Markdown-style content as styled paragraphs right after it.
We respect the document's heading hierarchy: 'Heading 2' for survey label,
'Heading 3' for question groups, 'Heading 4' for individual questions, and
List Bullet for sub-items / answer codes.

Inputs (relative to repository root):
  data/monograph_in.docx                  — monograph to be patched (rename your draft to this)
  results/Appendix_A_Dataset_A_questionnaire.md  — produced by 09_parse_lss.py
  results/Appendix_A_Dataset_B_questionnaire.md  — produced by 09_parse_lss.py
Output:
  results/monograph_with_appendix_a.docx
"""
from copy import deepcopy
from pathlib import Path
from docx import Document

ROOT = Path(__file__).resolve().parents[1]
SRC = str(ROOT / "data" / "monograph_in.docx")
DST = str(ROOT / "results" / "monograph_with_appendix_a.docx")
DATASET_A_FILE = ROOT / "results" / "Appendix_A_Dataset_A_questionnaire.md"
DATASET_B_FILE = ROOT / "results" / "Appendix_A_Dataset_B_questionnaire.md"


def parse_md_to_blocks(text):
    """Convert our Markdown render into a list of (style, text) tuples for python-docx."""
    blocks = []
    for line in text.splitlines():
        s = line.rstrip()
        if not s.strip():
            blocks.append(("Normal", ""))
            continue
        if s.startswith("## "):
            blocks.append(("Heading 3", s[3:].strip()))
        elif s.startswith("### "):
            blocks.append(("Heading 4", s[4:].strip()))
        elif s.startswith("#### "):
            blocks.append(("Heading 5", s[5:].strip()))
        elif s.startswith("# "):
            blocks.append(("Heading 2", s[2:].strip()))
        elif s.lstrip().startswith("- "):
            blocks.append(("List Bullet", s.lstrip()[2:].strip()))
        elif s.lstrip().startswith("**") and s.rstrip().endswith("**"):
            blocks.append(("Normal_bold", s.strip().strip("*").strip()))
        elif s.startswith("Príloha:"):
            blocks.append(("Heading 2", s.strip()))
        else:
            blocks.append(("Normal", s.strip()))
    return blocks


def main():
    doc = Document(SRC)
    # Find the 'Normal: Appendix A' anchor: paragraph whose text starts with 'Appendix A'
    anchor_idx = None
    for i, p in enumerate(doc.paragraphs):
        if p.text.strip() == "Appendix A":
            anchor_idx = i
            break
    if anchor_idx is None:
        # Try alternative: look for "Appendix A" as substring
        for i, p in enumerate(doc.paragraphs):
            if p.text.strip().startswith("Appendix A"):
                anchor_idx = i
                break
    if anchor_idx is None:
        raise SystemExit("Anchor paragraph 'Appendix A' not found.")

    print(f"Anchor: paragraph #{anchor_idx} = '{doc.paragraphs[anchor_idx].text!r}'")

    anchor = doc.paragraphs[anchor_idx]

    # Ensure available styles; otherwise fall back to 'Normal'
    available_styles = {s.name for s in doc.styles}

    def add_after(prev_para, style_name, text):
        """Insert a new paragraph with given style right after prev_para; return new paragraph."""
        # Use a deepcopy of an existing paragraph element to inherit base XML structure
        new_el = deepcopy(prev_para._element)
        # Strip existing runs
        for r in new_el.findall(".//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}r"):
            r.getparent().remove(r)
        prev_para._element.addnext(new_el)
        # Re-fetch as a paragraph object
        from docx.text.paragraph import Paragraph
        new_para = Paragraph(new_el, prev_para._parent)
        # Apply style
        target = style_name
        if style_name == "Normal_bold":
            target = "Normal"
        if target in available_styles:
            new_para.style = doc.styles[target]
        else:
            # Fallback to Normal
            if "Normal" in available_styles:
                new_para.style = doc.styles["Normal"]
        run = new_para.add_run(text)
        if style_name == "Normal_bold":
            run.bold = True
        return new_para

    # Read both files
    md_a = DATASET_A_FILE.read_text(encoding="utf-8")
    md_b = DATASET_B_FILE.read_text(encoding="utf-8")

    blocks = []
    blocks.append(("Heading 2", "Appendix A — Plné znenie dotazníkov"))
    blocks.append(("Normal", "Táto príloha obsahuje úplnú štruktúru dotazníkov, ktoré boli použité na zber Dataset A (digitálne návyky) a Dataset B (digitálne prostredie a smart-city gramotnosť). Otázky sú zoskupené podľa pôvodných LimeSurvey skupín (groups). Pre každú otázku sa uvádza jej kód, typ, znenie, prípadné položky (sub-questions) a možnosti odpovedí. Voľné textové polia s našepkávačmi (zoznam stredných škôl, zoznam slovenských miest) nie sú reprodukované v tlačovej prílohe — sú dostupné v zdrojovom LimeSurvey exporte v replikačnom balíku."))
    blocks.append(("Normal", ""))
    blocks.append(("Heading 3", "A.1 Dataset A — Digitálne návyky a používanie mobilných aplikácií"))
    # Skip the first line of dataset A md (it duplicates the title)
    md_a_body = md_a.split("\n", 1)[1] if "\n" in md_a else md_a
    blocks.extend(parse_md_to_blocks(md_a_body))
    blocks.append(("Normal", ""))
    blocks.append(("Heading 3", "A.2 Dataset B — Digitálne prostredie a smart-city gramotnosť"))
    md_b_body = md_b.split("\n", 1)[1] if "\n" in md_b else md_b
    blocks.extend(parse_md_to_blocks(md_b_body))

    # Insert in reverse order? No — insert sequentially after anchor moving cursor each step
    cursor = anchor
    for style_name, text in blocks:
        cursor = add_after(cursor, style_name, text)

    doc.save(DST)
    print(f"Saved to {DST} — {len(blocks)} paragraphs added")


if __name__ == "__main__":
    main()
