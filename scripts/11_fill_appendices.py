"""Fill Appendices B, C, D, E in the monograph with real content from the
public aggregate CSVs and from the LSS-derived codebook information.

Strategy:
- Locate paragraphs 'Appendix B', 'Appendix C', 'Appendix D', 'Appendix E'.
- After each anchor, REPLACE the existing skeleton-only "Table B1.1 ..." lines
  with proper Word tables + descriptive paragraphs. To stay safe with python-docx,
  we leave the anchor and its skeleton placeholders, but immediately after each
  Appendix header insert the new content as a sequence of styled paragraphs
  + native Word tables.
- Skeleton placeholder lines that remain redundant after our insertion are
  removed at the end.

Inputs (relative to repository root):
  data/monograph_in.docx           — monograph to be patched
  data/aggregates/*.csv            — produced by 08_export_aggregates.py
Output:
  results/monograph_with_appendices.docx
"""
from copy import deepcopy
from pathlib import Path
import csv
from docx import Document
from docx.shared import Pt
from docx.text.paragraph import Paragraph

ROOT = Path(__file__).resolve().parents[1]
SRC = str(ROOT / "data" / "monograph_in.docx")
DST = str(ROOT / "results" / "monograph_with_appendices.docx")
AGG = ROOT / "data" / "aggregates"

W_NS = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'


def read_csv(name):
    with open(AGG / name, encoding='utf-8') as f:
        rows = list(csv.reader(f))
    return rows


def add_paragraph_after(prev_para, doc, style='Normal', text='', bold=False):
    new_el = deepcopy(prev_para._element)
    for r in new_el.findall(f".//{W_NS}r"):
        r.getparent().remove(r)
    prev_para._element.addnext(new_el)
    new_para = Paragraph(new_el, prev_para._parent)
    available = {s.name for s in doc.styles}
    if style in available:
        new_para.style = doc.styles[style]
    else:
        new_para.style = doc.styles['Normal']
    if text:
        run = new_para.add_run(text)
        if bold:
            run.bold = True
    return new_para


def add_table_after(prev_para, doc, header, rows, col_widths=None):
    """Insert a Word table after prev_para. Returns the new table object and a
    paragraph object positioned after the table for further insertions."""
    # Create a placeholder paragraph after prev_para
    spacer = add_paragraph_after(prev_para, doc, 'Normal', '')
    # Add the table at end of document, then move it to position after spacer
    n_cols = len(header)
    table = doc.add_table(rows=1 + len(rows), cols=n_cols)
    try:
        table.style = 'Light List Accent 1'
    except Exception:
        try:
            table.style = 'Table Grid'
        except Exception:
            pass
    # Header
    for j, h in enumerate(header):
        cell = table.rows[0].cells[j]
        cell.text = ''
        run = cell.paragraphs[0].add_run(str(h))
        run.bold = True
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            cell = table.rows[1 + i].cells[j]
            cell.text = ''
            cell.paragraphs[0].add_run('' if val is None else str(val))

    # Move the table xml from end-of-body to right after spacer
    body = doc.element.body
    tbl_el = table._tbl
    body.remove(tbl_el)  # detach
    spacer._element.addnext(tbl_el)

    # Add an empty paragraph after the table so subsequent insertions can latch on
    after = doc.add_paragraph()
    after_el = after._element
    body.remove(after_el)
    tbl_el.addnext(after_el)
    after.style = doc.styles['Normal']
    return table, after


def find_paragraph(doc, exact_text):
    for i, p in enumerate(doc.paragraphs):
        if p.text.strip() == exact_text:
            return i, p
    return -1, None


# ---------------------------------------------------------------------------
def fill_appendix_b(doc, anchor):
    """Codebook overview for both surveys (variable -> label -> type -> coding)."""
    cur = anchor
    cur = add_paragraph_after(cur, doc, 'Heading 2', 'Appendix B — Codebook (Datasets A and B)')
    cur = add_paragraph_after(cur, doc, 'Normal',
        'This appendix lists the analytical variables derived from the two survey instruments '
        '(see Appendix A for the full questionnaire texts). Each entry shows the LimeSurvey item code, '
        'the analytical label used in the replication scripts, the variable type and the coding rule. '
        'The same definitions are used by scripts/01_build_indices.py in the replication package.')
    cur = add_paragraph_after(cur, doc, 'Normal', '')

    # B.1 Dataset A
    cur = add_paragraph_after(cur, doc, 'Heading 3', 'B.1 Dataset A — mobile habits')
    rows_A = [
        ['POHLAVIE', 'gender', 'categorical', '"Muž" (M) / "Žena" (Z); 1 missing case excluded'],
        ['VEK', 'age', 'integer', '15–17 in the analytical SS sample'],
        ['TYPSKOLY', 'school_type', 'categorical', 'Filter: only "Stredná škola" (SS) retained'],
        ['OS', 'os', 'categorical', '"iOS" / "Android"'],
        ['ZNACKA', 'phone_brand', 'free text', 'Descriptive only'],
        ['POCETAPLIK', 'inst (apps_installed)', 'integer', 'Self-reported number of installed apps; AEI input'],
        ['POCETAPLIKpouzivam', 'used (apps_used)', 'integer', 'Self-reported number of actively used apps; descriptive'],
        ['G02Q51', 'time_h', 'ordinal → midpoints', '<1h→0.5; 1–3h→2; 3–5h→4; >5h→6'],
        ['G02Q52', 'last_n (recency)', 'ordinal 1–3', '"Viac ako pred mesiacom"=1; "Tento mesiac"=2; "Tento týždeň"=3'],
        ['APP1NAZOV…APP3NAZOV', 'top1_name…top3_name', 'free text', 'Open name of most-used app'],
        ['APP1DOVOD…APP3DOVOD', 'top1_reason…top3_reason', 'free text', 'Open reason'],
        ['G01Q53[SQ001-007]', 'm_peer, m_family, m_ads, m_curiosity, m_school, m_entertain, m_trend', 'binary (0/1)', 'Multi-check motives for installing apps'],
        ['G01Q55[SQ001-007]', 'interest_*', 'binary (0/1)', 'Multi-check personal interests (max 3)'],
        ['G01Q56[SQ001-004]', 'adopter_*', 'binary (0/1)', 'Adopter posture statements'],
        ['G01Q57[SQ001-004]', 'edu_attitude_*', 'binary (0/1)', 'Attitudes towards educational/civic apps'],
        ['(derived)', 'use_ratio', 'float', 'apps_used ÷ apps_installed'],
        ['(derived)', 'motives_count', 'integer 0–7', 'Sum of m_* binary motives'],
        ['(derived)', 'AEI', '0–100', '0.5·minmax(time_h) + 0.3·minmax(inst) + 0.2·minmax(last_n)'],
    ]
    table, cur = add_table_after(cur, doc,
        ['LimeSurvey item', 'Analytical name', 'Type', 'Coding / values'],
        rows_A)
    cur = add_paragraph_after(cur, doc, 'Normal', '')

    # B.2 Dataset B
    cur = add_paragraph_after(cur, doc, 'Heading 3', 'B.2 Dataset B — digital environment & smart-city literacy')
    rows_B = [
        ['G01Q07', 'school_name', 'free text', 'School name (autocomplete; not analysed individually)'],
        ['G01Q08', 'is_student', 'binary', '"Áno"/"Nie"; only "Áno" rows retained'],
        ['G02Q02', 'grade_int', 'integer', '1–4 used; rare 5+ excluded as non-mainstream'],
        ['G01Q03', 'gender', 'categorical', '"Muž" / "Žena"'],
        ['G01Q20', 'city', 'free text', 'City of residence (autocomplete)'],
        ['G02Q04 / G02Q04Copy', 'edu_mother / edu_father', 'ordinal 1–3', 'základné=1; stredoškolské=2; vysokoškolské=3'],
        ['G02Q02Copy', 'siblings', 'integer', 'Number of siblings'],
        ['G03Q08', 'sc_definition_correct', 'binary', '1 if conceptual definition selected'],
        ['G01Q09', 'sc_not_part', 'categorical', 'Multiple choice'],
        ['G01Q10', 'sc_role_tech (csa_role)', 'binary', '1 if "improve aspects of urban life through data and connectivity" chosen'],
        ['G01Q22', 'sc_at_school (csa_school)', 'ordinal 0/1/2', 'Detailed=2; Briefly=1; Other=0'],
        ['G01Q23', 'sc_importance_edu (csa_imp)', 'Likert 1–5', 'Nepodstatné=1 … Veľmi dôležité=5'],
        ['G05Q18', 'encountered_local (csa_local)', 'binary', '"Áno"=1; "Nie"=0'],
        ['G01Q19[SQ001-021]', 'feature_*', 'binary (0/1)', '21 smart-city feature items'],
        ['G01Q26[SQ001]', 'q_school (school internet)', 'Likert 1–5', 'SDR component'],
        ['G01Q26[SQ002]', 'q_home (home internet)', 'Likert 1–5', 'DAE component'],
        ['G01Q26[SQ003]', 'q_city (city internet)', 'Likert 1–5', 'Descriptive only'],
        ['G01Q26[SQ004]', 'q_equip (school equipment)', 'Likert 1–5', 'SDR component'],
        ['G08Q25[SQ001-004]', 'dev_mobile/media/compute/iot', 'integer 0–30', 'Winsorised at 30; sum → dev_total; mean → dev_mean (DAE input)'],
        ['(derived)', 'DAE', '0–100', 'minmax(z(q_home) + z(dev_total))/2'],
        ['(derived)', 'SDR', '0–100', 'minmax(z(q_school) + z(q_equip))/2'],
        ['(derived)', 'CSA', '0–100', 'minmax(mean of z(csa_school, csa_role, csa_imp, csa_local))'],
        ['(derived)', 'YSCR', '0–100', '0.5·CSA + 0.3·SDR(mean-imputed) + 0.2·DAE(mean-imputed)'],
        ['(derived)', 'YSCR_listwise', '0–100', '0.5·CSA + 0.3·SDR + 0.2·DAE (listwise; sensitivity)'],
    ]
    table, cur = add_table_after(cur, doc,
        ['LimeSurvey item', 'Analytical name', 'Type', 'Coding / values'],
        rows_B)
    cur = add_paragraph_after(cur, doc, 'Normal', '')
    cur = add_paragraph_after(cur, doc, 'Normal',
        'Listwise deletion is used inside each first-order index. Headline valid N values throughout the monograph: '
        'Dataset A — total = 82, AEI = 77, apps_used = 80, time = 81. Dataset B — total = 401, mainstream grades 1–4 = 368, '
        'DAE valid = 342, SDR valid = 354 (or 341 if both items required), CSA valid = 368, Model 7.2 = 332. '
        'Self-reported device counts above 30 are treated as data-entry errors (winsorised to NaN); fewer than five cases are affected.')
    return cur


def fill_appendix_c(doc, anchor):
    """Step-by-step index construction tables."""
    cur = anchor
    cur = add_paragraph_after(cur, doc, 'Heading 2', 'Appendix C — Index Construction')
    cur = add_paragraph_after(cur, doc, 'Normal',
        'The five composite indicators follow the unified workflow described in §4.4: (a) input items are recoded to a numeric scale; '
        '(b) each item is z-scored across the analytical sample; (c) component z-scores are averaged into a raw index for each respondent; '
        '(d) the raw index is rescaled to 0–100 via min–max normalisation. Each table below specifies the exact inputs and operations used '
        'by scripts/01_build_indices.py to reproduce the published values.')
    cur = add_paragraph_after(cur, doc, 'Normal', '')

    # AEI
    cur = add_paragraph_after(cur, doc, 'Heading 3', 'Table C1 — App Engagement Index (AEI)')
    rows = [
        ['1', 'Inputs', 'time_h, inst (apps_installed), last_n (recency 1–3)'],
        ['2', 'Recoding', 'time_h: <1h→0.5, 1–3h→2, 3–5h→4, >5h→6  •  last_n: 1=>1m ago, 2=last month, 3=last week'],
        ['3', 'Component rescaling', 'Each input rescaled to 0–100 via min–max'],
        ['4', 'Aggregation', 'AEI = 0.5·minmax(time_h) + 0.3·minmax(inst) + 0.2·minmax(last_n)'],
        ['5', 'Validity', 'Listwise: respondents with all three components valid (N = 77)'],
    ]
    _, cur = add_table_after(cur, doc, ['Step', 'Operation', 'Details'], rows)

    cur = add_paragraph_after(cur, doc, 'Heading 3', 'Table C2 — Digital Access Environment (DAE)')
    rows = [
        ['1', 'Inputs', 'q_home (G01Q26[SQ002], 1–5), dev_total = sum(dev_mobile, dev_media, dev_compute, dev_iot)'],
        ['2', 'Outlier handling', 'Each device count winsorised to [0, 30] before summation'],
        ['3', 'z-scoring', 'z(q_home), z(dev_total) over the mainstream subsample'],
        ['4', 'Raw index', 'DAE_raw = mean(z(q_home), z(dev_total))'],
        ['5', 'Rescale', 'DAE = 100·(DAE_raw − min) / (max − min) → range 0–100'],
        ['6', 'Validity', 'Listwise N = 342 (both inputs valid)'],
    ]
    _, cur = add_table_after(cur, doc, ['Step', 'Operation', 'Details'], rows)

    cur = add_paragraph_after(cur, doc, 'Heading 3', 'Table C3 — School Digital Readiness (SDR)')
    rows = [
        ['1', 'Inputs', 'q_school (G01Q26[SQ001], 1–5), q_equip (G01Q26[SQ004], 1–5)'],
        ['2', 'z-scoring', 'z(q_school), z(q_equip)'],
        ['3', 'Raw index', 'SDR_raw = mean(z(q_school), z(q_equip))'],
        ['4', 'Rescale', 'SDR = min–max → 0–100'],
        ['5', 'Validity', 'Listwise N = 354'],
    ]
    _, cur = add_table_after(cur, doc, ['Step', 'Operation', 'Details'], rows)

    cur = add_paragraph_after(cur, doc, 'Heading 3', 'Table C4 — Curricular and Smart-City Awareness (CSA)')
    rows = [
        ['1', 'Inputs', 'csa_school (G01Q22, 0/1/2), csa_role (G01Q10, 0/1), csa_imp (G01Q23, 1–5), csa_local (G05Q18, 0/1)'],
        ['2', 'z-scoring', 'Four item-level z-scores'],
        ['3', 'Raw index', 'CSA_raw = mean of available z-scores; ≥3 of 4 components required'],
        ['4', 'Rescale', 'CSA = min–max → 0–100'],
        ['5', 'Validity', 'N = 368 (mainstream grades 1–4)'],
    ]
    _, cur = add_table_after(cur, doc, ['Step', 'Operation', 'Details'], rows)

    cur = add_paragraph_after(cur, doc, 'Heading 3', 'Table C5 — Youth Smart-City Readiness (YSCR)')
    rows = [
        ['1', 'Inputs', 'CSA, SDR, DAE (each 0–100)'],
        ['2', 'Default weights', '0.5 · CSA + 0.3 · SDR + 0.2 · DAE'],
        ['3', 'Missing-data rule (default)', 'Mean-imputation for missing SDR and DAE before applying weights → preserves N = 368'],
        ['4', 'Listwise sensitivity', 'YSCR_listwise = same formula without imputation; valid N = 332; Pearson r with default ≈ 1.000'],
        ['5', 'Alternative weighting', 'Equal weights (1/3 each) and PCA-derived weights (0.276 / 0.376 / 0.348) reported in §6.1; pairwise Pearson r ≥ 0.968'],
    ]
    _, cur = add_table_after(cur, doc, ['Step', 'Operation', 'Details'], rows)
    return cur


def fill_appendix_d(doc, anchor):
    """Descriptive statistics and correlations from CSV aggregates."""
    cur = anchor
    cur = add_paragraph_after(cur, doc, 'Heading 2', 'Appendix D — Descriptive Statistics and Correlations (Aggregated)')
    cur = add_paragraph_after(cur, doc, 'Normal',
        'All values below are reproduced from the public CC BY 4.0 aggregate CSV files in '
        'data/aggregates/ of the replication package (Zenodo DOI 10.5281/zenodo.20011751). '
        'No individual-level values are released.')
    cur = add_paragraph_after(cur, doc, 'Normal', '')

    # D1 — Dataset A descriptives
    cur = add_paragraph_after(cur, doc, 'Heading 3', 'Table D1 — Descriptive statistics, Dataset A')
    rows = read_csv('dataset_A_aggregates.csv')
    _, cur = add_table_after(cur, doc, rows[0], rows[1:])

    # D2 — Dataset B descriptives
    cur = add_paragraph_after(cur, doc, 'Heading 3', 'Table D2 — Descriptive statistics, Dataset B (incl. YSCR_listwise sensitivity)')
    rows = read_csv('dataset_B_aggregates.csv')
    _, cur = add_table_after(cur, doc, rows[0], rows[1:])

    # D3 — Pearson correlations
    cur = add_paragraph_after(cur, doc, 'Heading 3', 'Table D3 — Pearson correlations between core indices (Dataset B)')
    rows = read_csv('dataset_B_correlations.csv')
    header = ['Index'] + rows[0][1:]
    body = [r for r in rows[1:]]
    _, cur = add_table_after(cur, doc, header, body)

    # D4 — Bootstrap CIs
    cur = add_paragraph_after(cur, doc, 'Heading 3', 'Table D4 — Bootstrap 95 % confidence intervals (2 000 resamples)')
    rows = read_csv('bootstrap_ci.csv')
    _, cur = add_table_after(cur, doc, rows[0], rows[1:])

    # D5 — YSCR sensitivity descriptives
    cur = add_paragraph_after(cur, doc, 'Heading 3', 'Table D5 — YSCR sensitivity to weighting scheme (default vs equal vs PCA)')
    rows = read_csv('yscr_sensitivity_descriptives.csv')
    header = ['Statistic'] + rows[0][1:]
    body = [r for r in rows[1:]]
    _, cur = add_table_after(cur, doc, header, body)

    cur = add_paragraph_after(cur, doc, 'Heading 3', 'Table D6 — Pearson correlations among YSCR weighting variants')
    rows = read_csv('yscr_sensitivity_correlations.csv')
    header = ['Variant'] + rows[0][1:]
    body = [r for r in rows[1:]]
    _, cur = add_table_after(cur, doc, header, body)

    # D7 — by gender
    cur = add_paragraph_after(cur, doc, 'Heading 3', 'Table D7 — Index scores by gender (Dataset B)')
    rows = read_csv('dataset_B_by_gender.csv')
    flat_rows = []
    for r in rows[3:]:  # skip multi-line header
        flat_rows.append(r)
    header = ['Gender',
              'DAE M', 'DAE SD', 'DAE N',
              'SDR M', 'SDR SD', 'SDR N',
              'CSA M', 'CSA SD', 'CSA N',
              'YSCR M', 'YSCR SD', 'YSCR N']
    _, cur = add_table_after(cur, doc, header, flat_rows)

    # D8 — by grade
    cur = add_paragraph_after(cur, doc, 'Heading 3', 'Table D8 — Index scores by grade (Dataset B)')
    rows = read_csv('dataset_B_by_grade.csv')
    flat_rows = []
    for r in rows[3:]:
        flat_rows.append(r)
    header = ['Grade',
              'DAE M', 'DAE SD', 'DAE N',
              'SDR M', 'SDR SD', 'SDR N',
              'CSA M', 'CSA SD', 'CSA N',
              'YSCR M', 'YSCR SD', 'YSCR N']
    _, cur = add_table_after(cur, doc, header, flat_rows)

    # D9 — Dataset A motives
    cur = add_paragraph_after(cur, doc, 'Heading 3', 'Table D9 — Distribution of installation motives, Dataset A')
    rows = read_csv('dataset_A_motives.csv')
    _, cur = add_table_after(cur, doc, rows[0], rows[1:])
    return cur


def fill_appendix_e(doc, anchor):
    """Cluster diagnostics + segment profiles."""
    cur = anchor
    cur = add_paragraph_after(cur, doc, 'Heading 2', 'Appendix E — Youth Segments (Cluster diagnostics and profiles)')
    cur = add_paragraph_after(cur, doc, 'Normal',
        'Segments derived from k-means clustering on standardised behavioural and motivational features of Dataset A '
        '(features: AEI, motives_count, m_entertain, m_school, m_trend, m_curiosity; random_state = 42, n_init = 20). '
        'Solutions for k = 2..6 are reported as diagnostic context; k = 4 is the solution interpreted in Chapter 8 and Table 40.')
    cur = add_paragraph_after(cur, doc, 'Normal', '')

    # E1 - diagnostics
    cur = add_paragraph_after(cur, doc, 'Heading 3', 'Table E1 — Cluster-solution diagnostics (k = 2..6)')
    rows = read_csv('cluster_diagnostics.csv')
    _, cur = add_table_after(cur, doc, rows[0], rows[1:])

    # E2 - profile
    cur = add_paragraph_after(cur, doc, 'Heading 3', 'Table E2 — Cluster centroid profile (k = 4)')
    rows = read_csv('cluster_profiles_k4.csv')
    _, cur = add_table_after(cur, doc, rows[0], rows[1:])

    # E3 - persona narrative cross-walk
    cur = add_paragraph_after(cur, doc, 'Heading 3', 'Table E3 — Cluster → persona crosswalk')
    persona_rows = [
        ['0', '24 (31.2 %)', '40.95', 'Pragmatic utilitarians', 'Moderate AEI; broad motive set including school/work and curiosity; balanced app sets'],
        ['1', '25 (32.5 %)', '57.05', 'Social-entertainment heavy users', 'High AEI; narrow motive set dominated by entertainment; frequent social/streaming use'],
        ['2', '17 (22.1 %)', '37.74', 'Low-engagement / reluctant', 'Lower AEI; low novelty-seeking; narrow motive list'],
        ['3', '11 (14.3 %)', '63.02', 'Trend followers / early adopters', 'Highest AEI; multi-motive (entertainment, curiosity, trends); responsive to advertising'],
    ]
    _, cur = add_table_after(cur, doc, ['Cluster', 'N (%)', 'Mean AEI', 'Persona label', 'Profile'], persona_rows)

    cur = add_paragraph_after(cur, doc, 'Normal',
        'The personas are interpretive constructs (see caveat in Chapter 8). Reference CSA / YSCR ranges in Table 40 '
        'are illustrative anchors taken from the corresponding readiness segments in Dataset B; Datasets A and B '
        'are not merged at the individual level.')
    return cur


def remove_skeleton_after(doc, header_text, until_next_appendix=True):
    """After we have appended new content right after the appendix anchor,
    the original skeleton placeholder lines (Table B1.1, B1.2, …) sit further
    down in the document. We detect them and remove. We stop at the next
    Appendix header or the bibliography boundary."""
    in_skeleton = False
    to_remove = []
    skeleton_markers = (
        'Table B1.', 'Table B2.', 'Table C1', 'Table C2', 'Table C3',
        'Table C4', 'Table C5', 'Table D1', 'Table D2', 'Table E1',
        'B1. Codebook', 'B2. Codebook',
        '(kompaktný blok', '(N sú uvedené', '(Dataset A', '(Dataset B',
    )
    # Find the original-skeleton starting point: a paragraph equal to header_text
    # that is NOT a Heading 2 (we just inserted a Heading 2 with the same name)
    candidate_idxs = []
    for i, p in enumerate(doc.paragraphs):
        if p.text.strip() == header_text and p.style.name != 'Heading 2':
            candidate_idxs.append(i)
    for idx in candidate_idxs:
        # remove from idx until we hit either the next appendix header or end
        i = idx
        while i < len(doc.paragraphs):
            p = doc.paragraphs[i]
            t = p.text.strip()
            # stop at next Appendix anchor or empty-section boundary
            if i != idx and (t.startswith('Appendix ') and t != header_text):
                break
            # stop at "Bibliography" if we ever go that far
            if t == 'Bibliography':
                break
            # only remove the original skeleton stuff (the duplicated anchor itself + the skeletal table captions)
            if i == idx or t == '' or any(t.startswith(m) for m in skeleton_markers):
                to_remove.append(p)
                i += 1
                continue
            else:
                # Hit a non-skeleton line — stop
                break
    for p in to_remove:
        try:
            p._element.getparent().remove(p._element)
        except Exception:
            pass


def main():
    doc = Document(SRC)
    print(f'Loaded: {len(doc.paragraphs)} paragraphs, {len(doc.tables)} tables')

    # Anchors are existing skeleton "Appendix X" Normal paragraphs. We'll insert
    # the new content right BEFORE those anchors, then strip the anchors+skeleton.

    # Process E first, then D, C, B (so insertions don't shift earlier anchor indices)
    # Actually: each anchor lookup is dynamic; safer is to do them in document order
    # but use add_paragraph_after on the anchor *paragraph object*; modifications further
    # down don't affect the saved object reference.

    for anchor_text, fill_func in [
        ('Appendix B – Codebook (Datasets A and B)', fill_appendix_b),
        ('Appendix C – Index Construction', fill_appendix_c),
        ('Appendix D – Descriptive Statistics and Correlations (Aggregated)', fill_appendix_d),
        ('Appendix E – Youth Segments (Template)', fill_appendix_e),
    ]:
        idx, p = find_paragraph(doc, anchor_text)
        if idx < 0:
            print(f'  ANCHOR NOT FOUND: {anchor_text}')
            continue
        print(f'  Filling {anchor_text} (anchor at #{idx})')
        fill_func(doc, p)
        # Now remove the original anchor + its skeleton placeholders (they sit AFTER our inserted content)
        remove_skeleton_after(doc, anchor_text)

    doc.save(DST)
    print(f'Saved {DST}')


if __name__ == '__main__':
    main()
